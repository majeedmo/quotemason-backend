# /// script
# requires-python = ">=3.10"
# dependencies = ["pypdf>=4.0", "cryptography>=3.1", "python-docx>=1.1"]
# ///
"""Produce the redacted past-quote corpus from the original client quotes.

Run:  uv run scripts/redact_quotes.py

Reads quotes/*.docx|pdf and quotes/redaction-map.json (both LOCAL ONLY — the
quotes/ folder is gitignored and must never be committed), removes client PII
and the contractor's identity, and writes metadata-tagged markdown to
corpus/quotes-redacted/ for RAG ingestion.

This script is deliberately PII-free: every name, address, filename, and
business identifier lives in quotes/redaction-map.json. When a new quote
arrives, register it there (file entry + the client's name/street tokens)
and re-run.

Address policy (2026-07-12): the house NUMBER is anonymized to "[NO]" but the
street name and city are KEPT — street + city is the zoning-lookup signal.
The contractor's own address is the exception: it identifies the business and
is removed entirely.

Redacted: client names/emails/phones, house numbers, postal codes,
permit/estimate numbers, contractor name (-> "Company A") and address and
licence/HST/WSIB identifiers. Package tiers, scope vocabulary, streets,
cities, line items, allowances, prices, and timelines are preserved — they
are the estimation and zoning signal.

After writing, the script greps its own output for leakage and reports it.
"""

from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from docx import Document
from pypdf import PdfReader

REPO = Path(__file__).resolve().parent.parent
SRC = REPO / "quotes"
OUT = REPO / "corpus/quotes-redacted"
MAP = SRC / "redaction-map.json"


# Generic-fallback street types: unambiguous suffixes only. Ambiguous words that
# appear in construction text (Walk, Gate, Green, Common, Row, Close, Hill, Ridge,
# Landing, Mews, Heights, Grove) are matched only via the map's explicit street
# lists — add such streets to quotes/redaction-map.json when they appear.
STREET_TYPES = (
    r"Road|Rd|Drive|Dr|Way|Place|Pl|Court|Crt|Avenue|Ave|Blvd|Boulevard|Street|St"
    r"|Lane|Ln|Crescent|Cres|Path|Trail|Terrace|Circle|Cir"
)


def build_rules(m: dict) -> list[tuple[re.Pattern, str]]:
    proj_streets = "|".join(re.escape(s) for s in m["project_street_names"])
    biz_streets = "|".join(re.escape(s) for s in m["business_street_names"])
    names = "|".join(m["client_name_tokens"])
    rules: list[tuple[re.Pattern, str]] = []
    rules.append((re.compile("|".join(m["business_name_patterns"])), m["business_replacement"]))
    rules.append((re.compile("|".join(m["business_domain_patterns"]), re.I), "[REDACTED_DOMAIN]"))
    rules.append((re.compile("|".join(m["business_id_patterns"])), "[REDACTED_ID]"))
    rules.append((re.compile("|".join(m["estimate_no_patterns"])), "[ESTIMATE_NO]"))
    rules.append((re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+"), "[EMAIL]"))
    rules.append((re.compile(r"\b\d{3}[-.\s]\d{3}[-.\s]\d{4}\b"), "[PHONE]"))
    rules.append((re.compile(r"\b[A-Z]\d[A-Z]\s?\d[A-Z]\d\b"), "[POSTAL_CODE]"))
    # contractor's own address identifies the business -> removed entirely
    rules.append((re.compile(r"(?:\b\d{1,5}[A-Za-z]?\s+)?(?:" + biz_streets + r")\b\.?", re.I), "[CONTRACTOR_ADDRESS]"))
    # policy: keep street name + city (zoning signal); anonymize the house number only
    rules.append((re.compile(r"\b\d{1,5}[A-Za-z]?\s+(" + proj_streets + r")\b", re.I), r"[NO] \1"))
    # generic fallback for streets not in the map: same treatment
    rules.append((re.compile(
        r"\b\d{1,5}[A-Za-z]?\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)?\s+(?:" + STREET_TYPES + r")\b\.?)(?!\w)",
        re.I), r"[NO] \1"))
    rules.append((re.compile(r"Permit\s+(?:BP|No\.?|#)\s*[\w-]+\s+[\d-]+", re.I), "Permit [PERMIT_NO]"))
    rules.append((re.compile(r"\b(?:" + names + r")\b", re.I), "[CLIENT]"))
    rules.append((re.compile(r"\[CLIENT\](?:[\s,&+-]+\[CLIENT\])+"), "[CLIENT]"))
    return rules


def build_leak_patterns(m: dict) -> list[re.Pattern]:
    tokens = m["client_name_tokens"] + m["business_street_names"]
    return [
        re.compile("|".join(m["business_name_patterns"] + m["business_domain_patterns"]), re.I),
        re.compile(r"[\w.+-]+@[\w-]+\.[\w.]{2,}"),
        re.compile(r"\b\d{3}[-.]\d{3}[-.]\d{4}\b"),
        re.compile(r"\b(?:" + "|".join(re.escape(t) for t in tokens) + r")\b", re.I),
        # a number immediately before a known project street = unredacted house number
        re.compile(r"\b\d{1,5}[A-Za-z]?\s+(?:" + "|".join(re.escape(s) for s in m["project_street_names"]) + r")\b", re.I),
    ]


# --- boilerplate trim ---------------------------------------------------------
#
# The source quotes carry contractor liability/marketing/administrative-policy
# boilerplate that is verbatim (or near-verbatim) identical across every quote
# it appears in (verified against all 23 corpus files before writing these
# patterns) and carries zero comparable-project signal — it dilutes the RAG
# corpus with redundant text that is, in any case, already the guideline
# doc's job (§5 quoting rules: deposit, milestones, change-order fees,
# warranty terms — injected verbatim into every draft). Trimmed here, at
# generation time, rather than patching the already-redacted output, so
# future quotes redacted through this script get the same treatment.
#
# Deliberately NOT touched: EXCLUSIONS (templated in most quotes but genuinely
# project-specific in at least one — home_addition — so a blanket cut risks
# losing real signal); PROJECT COST / RENOVATION-CONSTRUCTION MILESTONES
# (genuine per-project dollar totals and payment schedules); OUT OF SCOPE,
# ADD-ONS, NOTE (not surveyed); all numbered work-category line items.

# Source docx paragraphs carry soft line-wraps mid-sentence (extract_docx
# preserves them as literal \n), so boilerplate phrases must match across
# whitespace runs, not literal single spaces. Escaping word-by-word (rather
# than the whole phrase then substituting) avoids re.escape's own backslash
# before each space colliding with the \s+ substitution.
def _ws(phrase: str) -> str:
    return r"\s+".join(re.escape(tok) for tok in phrase.split())


# The unlabeled intro's liability paragraph — two middle-clause wordings seen
# ("the estimate and the specifications" vs "the "scope of work" below").
_INTRO_LIABILITY = re.compile(
    _ws("Company A herein after called the") + r"\s*[“\"]Contractor[”\"]\s*" +
    _ws("will provide labor and materials for the work as outlined in") + r"\s+"
    r"(?:" + _ws("the estimate and the specifications") + "|"
    r"the\s+[“\"]scope\s+of\s+work[”\"]\s+below" + r")\." +
    r".*?" + _ws("will be billed accordingly") + r"\.\s*",
    re.S)
# Permit-fee-responsibility clause — only the common "Town/City of X Permit"
# wording (the Ontario-Building-Code service-list variant seen in a couple of
# files is structurally different and left alone rather than risk a bad cut).
_INTRO_PERMIT = re.compile(
    _ws("All work will be done as per") + r"\s+(?:Town|[Cc]ity)\s+of\s+\w+\s+" +
    _ws("Permit for") + r".*?" + _ws("Fire rated Windows/Shutters etc") + r"\.\s*",
    re.S)
_INTRO_HOURS = re.compile(
    _ws("Work shall be between the hours of 7:30 AM to 7:30 PM. Should Contractor require "
        "occasional work outside these times, approval from the home owner should be obtained "
        "prior to the time that work is required.") + r"\s*")
_INTRO_NDA = re.compile(
    _ws("By Signing this contract Client is agreeing to get into an NDA in which he/she not "
        "allowed to deal directly with any sub-contractors until next two years provided by "
        "the Contractor.") + r"\s*")

# The full-page marketing bio ("WHO WE ARE?") through the page-footer
# contact/address block that consistently follows it.
_WHO_WE_ARE = re.compile(
    r"WHO WE ARE\??.*?\[CONTRACTOR_ADDRESS\][^\n]*\n"
    r"(?:Contact:[^\n]*\n)?(?:\[REDACTED_DOMAIN\][^\n]*\n)?",
    re.S)

# Named sections that are administrative/contractual boilerplate, wholesale —
# stripped from the heading line through to the next recognized section
# boundary (mirrors, without importing, the heading vocabulary
# app/ingestion/chunking.py's chunk_quote() already treats as a boundary).
_BOILERPLATE_SECTION_HEADS = re.compile(
    r"^(WARRANTY|AGREEMENT OF SERVICES|CHANGE ORDER POLICY|CHANGE ORDERS)\s*:?\s*$", re.M)
_NEXT_SECTION_BOUNDARY = re.compile(
    r"^\s*(?:\d{1,2}\s*\|?\s+[A-Z][A-Z0-9 /&()+.,'’-]{4,80}?:"  # numbered work category
    r"|SCOPE OF (?:PROJECT|WORK)"
    r"|EXCLUSIONS|OUT OF SCOPE|ADD[- ]?ONS?|PROJECT COST|NOTE:?"
    r"|RENOVATION\s*/\s*CONSTRUCTION MILESTONES"
    r"|AGREEMENT OF SERVICES|CHANGE ORDER(?:S)?(?:\s+POLICY)?|WARRANTY)\b",
    re.M)


def _strip_boilerplate_sections(text: str) -> str:
    out, pos = [], 0
    for m in _BOILERPLATE_SECTION_HEADS.finditer(text):
        if m.start() < pos:
            continue  # already consumed by a previous removal in this pass
        out.append(text[pos:m.start()])
        nxt = _NEXT_SECTION_BOUNDARY.search(text, m.end())
        pos = nxt.start() if nxt else len(text)
    out.append(text[pos:])
    return "".join(out)


def strip_boilerplate(text: str) -> str:
    text = _WHO_WE_ARE.sub("", text)
    text = _INTRO_LIABILITY.sub("", text)
    text = _INTRO_PERMIT.sub("", text)
    text = _INTRO_HOURS.sub("", text)
    text = _INTRO_NDA.sub("", text)
    text = _strip_boilerplate_sections(text)
    return text


def extract_docx(path: Path) -> str:
    d = Document(path)
    out = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for tb in d.tables:
        out.append("")
        for row in tb.rows:
            out.append(" | ".join(c.text.strip().replace("\n", " / ") for c in row.cells))
    return "\n".join(out)


def extract_pdf(path: Path) -> str:
    r = PdfReader(path)
    return "\n".join((p.extract_text() or "") for p in r.pages)


def main() -> None:
    m = json.loads(MAP.read_text())
    rules = build_rules(m)
    leak_patterns = build_leak_patterns(m)
    OUT.mkdir(parents=True, exist_ok=True)
    seen_hashes: dict[str, str] = {}
    written, leaks = 0, []
    for name in sorted(m["files"]):
        src = SRC / name
        if not src.exists():
            print(f"MISSING (skipped): {name}")
            continue
        digest = hashlib.sha256(src.read_bytes()).hexdigest()
        if digest in seen_hashes:
            print(f"duplicate content (skipped): -> same as {m['files'][seen_hashes[digest]]['code']}")
            continue
        seen_hashes[digest] = name
        meta = m["files"][name]
        raw = extract_docx(src) if src.suffix == ".docx" else extract_pdf(src)
        clean = raw
        for pat, repl in rules:
            clean = pat.sub(repl, clean)
        # boilerplate trim runs after PII redaction: its patterns reference
        # redaction output tokens like [CONTRACTOR_ADDRESS]
        clean = strip_boilerplate(clean)
        slug = f"{meta['code']}-{meta['city'].lower()}-{meta['tier'].lower()}-{meta['scope']}" + (
            "-revised" if meta["revised"] else "")
        fm = "\n".join([
            "---",
            f"project_code: '{meta['code']}'",
            "doc_type: 'past_project_quote'",
            "jurisdiction: 'ontario'",
            f"city: '{meta['city']}'",
            f"street: '{meta.get('street', 'unspecified')}'",
            f"package_tier: '{meta['tier']}'",
            f"scope: '{meta['scope']}'",
            f"revised: {str(meta['revised']).lower()}",
            "source_version: 'redacted 2026-07-12; original is local-only (quotes/ is gitignored)'",
            "---",
            "",
        ])
        out_path = OUT / f"{slug}.md"
        out_path.write_text(fm + clean + "\n")
        written += 1
        for pat in leak_patterns:
            for hit in pat.findall(clean):
                leaks.append((out_path.name, str(hit)[:60]))
        print(f"wrote {out_path.relative_to(REPO)}  ({len(clean):,} chars)")
    print(f"\n{written} files written to {OUT.relative_to(REPO)}")
    if leaks:
        print(f"\nPOSSIBLE LEAKS ({len(leaks)}):")
        for fname, hit in leaks[:40]:
            print(f"  {fname}: {hit}")
    else:
        print("leak scan: clean")


if __name__ == "__main__":
    main()
